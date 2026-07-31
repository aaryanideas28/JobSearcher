# File: src/utils/docx_compiler.py
"""Word Document compiler for resume and cover letter generation."""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Any, Literal
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.workflow.state import AgentState

WORKSPACE_ROOT = Path("storage_workspace")
TEMPLATE_DIR = WORKSPACE_ROOT / "templates"
GENERATED_DIR = WORKSPACE_ROOT / "generated"

TEMPLATE_STYLES: dict[str, dict[str, Any]] = {
    "minimal_ats": {
        "font_family": "Times New Roman",
        "accent_color": RGBColor(0, 0, 0),
        "accent_color_hex": "000000",
        "header_align": 1,  # Center
        "margins": 0.5,
        "space_before": 8,
        "space_after": 2,
    },
    "modern_tech": {
        "font_family": "Times New Roman",
        "accent_color": RGBColor(0, 0, 0),
        "accent_color_hex": "000000",
        "header_align": 1,  # Center
        "margins": 0.5,
        "space_before": 8,
        "space_after": 2,
    },
    "classic_executive": {
        "font_family": "Times New Roman",
        "accent_color": RGBColor(0, 0, 0),
        "accent_color_hex": "000000",
        "header_align": 1,  # Center
        "margins": 0.5,
        "space_before": 8,
        "space_after": 2,
    },
    "compact_onepage": {
        "font_family": "Times New Roman",
        "accent_color": RGBColor(0, 0, 0),
        "accent_color_hex": "000000",
        "header_align": 1,  # Center
        "margins": 0.5,
        "space_before": 8,
        "space_after": 2,
    }
}


def add_markdown_paragraph_runs(p: docx.Paragraph, text: str, font_name: str, font_size_pt: float, base_color = None):
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            val = part[2:-2]
            run = p.add_run(val)
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        if base_color:
            run.font.color.rgb = base_color


def fix_concatenated_skills(skill_str: str) -> str:
    """Fix common concatenated technical skills lists by splitting them with commas."""
    import re
    # Direct replacements for exact user examples (and variations)
    replacements = {
        "Go (Golang)Python SQL": "Go (Golang), Python, SQL",
        "Go(Golang)Python SQL": "Go (Golang), Python, SQL",
        "PostgreSQLRedis MongoDB": "PostgreSQL, Redis, MongoDB",
        "KubernetesTerraformCI/CD": "Kubernetes, Terraform, CI/CD",
        "Event-DrivenKafka System Design": "Event-Driven Architecture, Kafka, System Design",
        "Event-Driven Kafka System Design": "Event-Driven Architecture, Kafka, System Design",
    }
    for k, v in replacements.items():
        if k in skill_str:
            skill_str = skill_str.replace(k, v)
            
    # Glued keyword splitters (case insensitive regex replacements)
    glued_pairs = [
        (r'PostgreSQLRedis', 'PostgreSQL, Redis'),
        (r'RedisMongoDB', 'Redis, MongoDB'),
        (r'KubernetesTerraform', 'Kubernetes, Terraform'),
        (r'TerraformCI/CD', 'Terraform, CI/CD'),
        (r'Golang\)Python', 'Golang), Python'),
        (r'PythonSQL', 'Python, SQL'),
        (r'Event-DrivenKafka', 'Event-Driven Architecture, Kafka'),
        (r'KafkaSystem', 'Kafka, System'),
    ]
    for pattern, replacement in glued_pairs:
        skill_str = re.sub(pattern, replacement, skill_str, flags=re.IGNORECASE)
        
    # Also separate any capital letter preceded by a lowercase letter/parenthesis if not already separated
    # e.g., "DockerKubernetes" -> "Docker, Kubernetes"
    skill_str = re.sub(r'(?<=[a-z0-9)])(?=[A-Z])', ', ', skill_str)
    
    # Restore standard tech names that got split by lookaround
    skill_str = skill_str.replace("Postgre, SQL", "PostgreSQL")
    skill_str = skill_str.replace("Postgre, Sql", "PostgreSQL")
    skill_str = skill_str.replace("Mongo, DB", "MongoDB")
    skill_str = skill_str.replace("Mongo, Db", "MongoDB")
    
    # Clean up multiple commas, spaces
    skill_str = re.sub(r'\s*,\s*', ', ', skill_str)
    skill_str = re.sub(r'\s+', ' ', skill_str).strip()
    return skill_str


class DocxCompiler:
    """Render documents directly into Microsoft Word docx format inside the workspace."""

    def __init__(self, template_dir: str | Path | None = None) -> None:
        self.template_dir = Path(template_dir) if template_dir else TEMPLATE_DIR

    def _setup_document(self, template_id: str = "minimal_ats") -> docx.Document:
        """Create a document with configured margins and base font style."""
        doc = docx.Document()
        style_config = TEMPLATE_STYLES.get(template_id, TEMPLATE_STYLES["minimal_ats"])
        margin_val = style_config["margins"]
        font_name = style_config["font_family"]
        
        for section in doc.sections:
            section.top_margin = Inches(margin_val)
            section.bottom_margin = Inches(margin_val)
            section.left_margin = Inches(margin_val)
            section.right_margin = Inches(margin_val)
            
        style = doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        
        return doc

    def _enforce_tight_spacing(self, doc: docx.Document, template_id: str = "minimal_ats") -> None:
        """Enforce strict tight paragraph and line spacing for all paragraphs in the document."""
        style_config = TEMPLATE_STYLES.get(template_id, TEMPLATE_STYLES["minimal_ats"])
        for idx, p in enumerate(doc.paragraphs):
            is_header = False
            if idx < 4:
                text_lower = p.text.lower()
                if any(x in text_lower for x in ["📞", "✉️", "📍", "🔗", "💻", "🌐", "|", "@"]):
                    is_header = True
                elif idx == 0 or idx == 1:
                    is_header = True
            
            if is_header:
                p.paragraph_format.line_spacing = 1.0
                continue
                
            style_name = p.style.name if p.style else ""
            
            # Check if this paragraph is a Section Heading
            is_heading = False
            pPr = p._p.pPr
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None and pBdr.find(qn('w:bottom')) is not None:
                    is_heading = True
            
            if is_heading:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
            elif style_name == "List Bullet":
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0
            elif p.paragraph_format.tab_stops:
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.line_spacing = 1.0
            else:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.line_spacing = 1.0

    def _add_section_title(self, doc: docx.Document, title: str, template_id: str = "minimal_ats"):
        """Add uppercase section title with horizontal divider line."""
        style_config = TEMPLATE_STYLES.get(template_id, TEMPLATE_STYLES["minimal_ats"])
        font_name = style_config["font_family"]
        accent_color = style_config["accent_color"]
        accent_hex = style_config["accent_color_hex"]

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(title.upper())
        run.bold = True
        run.font.name = font_name
        run.font.size = Pt(11)
        run.font.color.rgb = accent_color

        # XML Paragraph bottom border for perfect solid lines using template color
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # 6/8 pt = 0.75 pt
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), accent_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def parse_resume_text_to_dict(self, text: str) -> dict[str, Any]:
        """Parse raw resume text (with headings and divider lines) into a structured dictionary."""
        data = {
            "contact": {
                "name": "Candidate",
                "email": "",
                "phone": "",
                "location": "",
                "linkedin": "",
                "github": "",
                "portfolio": ""
            },
            "subtitle": "",
            "summary": "",
            "experience": [],
            "skills": [],
            "education": [],
            "projects": [],
            "certifications": [],
            "achievements": []
        }

        if not text:
            return data

        lines = [line.strip() for line in text.splitlines()]

        # 1. Parse header (up to the first section header or divider)
        header_lines = []
        section_start_idx = 0
        for i, line in enumerate(lines):
            is_div = line == "---" or re.match(r'^[-=*_]+$', line)
            is_hdr = any(s in line.upper() for s in ["EDUCATION", "EXPERIENCE", "PROJECTS", "SKILLS", "ACHIEVEMENTS", "SUMMARY", "EXTRACURRICULAR"]) and len(line) < 40
            if is_div or is_hdr:
                section_start_idx = i
                break
            if line:
                header_lines.append(line)

        if header_lines:
            data["contact"]["name"] = header_lines[0]
            for hl in header_lines[1:]:
                is_contact = False
                for k in ["phone", "email", "location", "github", "linkedin", "portfolio", "http", "https", "@", "|"]:
                    if k in hl.lower():
                        is_contact = True
                        break
                if re.search(r'\+?\d[\d\s\-()]{8,15}\d', hl):
                    is_contact = True

                if not is_contact:
                    if not data["subtitle"] and len(hl) < 60:
                        data["subtitle"] = hl
                        continue
                
                parts = [p.strip() for p in re.split(r'\||•|·', hl)]
                for p in parts:
                    if "@" in p:
                        m = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', p)
                        if m:
                            data["contact"]["email"] = m.group(0)
                    elif re.search(r'\+?\d[\d\s\-()]{8,15}\d', p):
                        m = re.search(r'(\+?\d[\d\s\-()]{8,15}\d)', p)
                        if m:
                            data["contact"]["phone"] = m.group(0)
                    elif any(domain in p.lower() for domain in ["linkedin.com", "linkedin", "in/"]):
                        data["contact"]["linkedin"] = p
                    elif any(domain in p.lower() for domain in ["github.com", "github"]):
                        data["contact"]["github"] = p
                    elif any(domain in p.lower() for domain in ["portfolio", "website", "http", "https"]):
                        data["contact"]["portfolio"] = p
                    elif len(p) > 3 and not data["contact"]["location"]:
                        data["contact"]["location"] = p

        # 2. Extract sections
        current_section = None
        section_content = []
        sections = {}

        for line in lines[section_start_idx:]:
            if line == "---" or re.match(r'^[-=*_]+$', line):
                continue
            
            line_upper = line.upper().replace("&", "AND").replace("AND EXTRACURRICULARS", "").replace("AND EXTRACURRICULAR", "")
            is_header = False
            header_name = None
            for s in ["EDUCATION", "EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "PROJECTS", "SKILLS", "ACHIEVEMENTS", "SUMMARY", "PROFILE SUMMARY", "EXTRACURRICULARS", "EXTRACURRICULAR", "ONLINE COURSES", "CERTIFICATIONS"]:
                if s in line_upper and len(line) < 40:
                    is_header = True
                    header_name = s
                    break

            if is_header:
                if current_section:
                    sections[current_section] = section_content
                current_section = header_name
                section_content = []
            else:
                if current_section and line:
                    section_content.append(line)

        if current_section:
            sections[current_section] = section_content

        # 3. Parse sections content
        # Profile Summary
        sum_lines = sections.get("SUMMARY") or sections.get("PROFILE SUMMARY") or []
        data["summary"] = " ".join(sum_lines)

        # Skills
        sk_lines = sections.get("SKILLS") or []
        for sl in sk_lines:
            if ":" in sl:
                parts = sl.split(":", 1)
                cat = parts[0].strip().lstrip("-*•· ").strip()
                items = [item.strip() for item in parts[1].split(",") if item.strip()]
                data["skills"].append({"category": cat, "items": items})
            else:
                cleaned = sl.lstrip("-*•· ").strip()
                if cleaned:
                    data["skills"].append(cleaned)

        # Experience
        exp_lines = sections.get("EXPERIENCE") or sections.get("PROFESSIONAL EXPERIENCE") or sections.get("WORK EXPERIENCE") or []
        current_entry = None
        for el in exp_lines:
            if el.startswith("•") or el.startswith("-") or el.startswith("*") or el.startswith("·"):
                if current_entry:
                    current_entry["bullets"].append(el.lstrip("-*•· ").strip())
            else:
                if current_entry:
                    data["experience"].append(current_entry)
                
                parts = [p.strip() for p in re.split(r'\||\t|  {3,}', el)]
                role = parts[0] if len(parts) > 0 else "Developer"
                company = parts[1] if len(parts) > 1 else ""
                location = parts[2] if len(parts) > 2 else ""
                dates = parts[3] if len(parts) > 3 else ""
                
                current_entry = {
                    "role": role,
                    "company": company,
                    "location": location,
                    "dates": dates,
                    "bullets": []
                }
        if current_entry:
            data["experience"].append(current_entry)

        # Projects
        proj_lines = sections.get("PROJECTS") or []
        current_proj = None
        for pl in proj_lines:
            if pl.startswith("•") or pl.startswith("-") or pl.startswith("*") or pl.startswith("·"):
                if current_proj:
                    current_proj["bullets"].append(pl.lstrip("-*•· ").strip())
            else:
                if current_proj:
                    data["projects"].append(current_proj)
                
                parts = [p.strip() for p in re.split(r'\||\t|  {3,}', pl)]
                name = parts[0] if len(parts) > 0 else "Project"
                tech = parts[1] if len(parts) > 1 else ""
                dates = parts[2] if len(parts) > 2 else ""
                
                current_proj = {
                    "name": name,
                    "technologies": tech,
                    "dates": dates,
                    "bullets": []
                }
        if current_proj:
            data["projects"].append(current_proj)

        # Education
        edu_lines = sections.get("EDUCATION") or []
        current_edu = None
        for edl in edu_lines:
            if any(k in edl.lower() for k in ["bachelor", "b.tech", "degree", "master", "m.tech", "12th", "10th", "school", "university", "technology"]):
                if current_edu:
                    data["education"].append(current_edu)
                
                parts = [p.strip() for p in re.split(r'\||\t|  {3,}', edl)]
                degree = parts[0]
                institution = parts[1] if len(parts) > 1 else ""
                location = parts[2] if len(parts) > 2 else ""
                dates = parts[3] if len(parts) > 3 else ""
                grade = parts[4] if len(parts) > 4 else ""
                
                current_edu = {
                    "degree": degree,
                    "institution": institution,
                    "location": location,
                    "dates": dates,
                    "grade": grade
                }
            elif current_edu:
                if not current_edu["institution"]:
                    current_edu["institution"] = edl
                elif not current_edu["dates"] and re.search(r'\d{4}', edl):
                    current_edu["dates"] = edl
                elif not current_edu["grade"] and any(k in edl.lower() for k in ["gpa", "cgpa", "%", "percent", "grade"]):
                    current_edu["grade"] = edl
                
        if current_edu:
            data["education"].append(current_edu)

        # Achievements
        ach_lines = sections.get("ACHIEVEMENTS") or sections.get("EXTRACURRICULARS") or sections.get("EXTRACURRICULAR") or []
        for al in ach_lines:
            cleaned = al.lstrip("-*•· ").strip()
            if cleaned:
                data["achievements"].append(cleaned)

        # Certifications
        cert_lines = sections.get("CERTIFICATIONS") or sections.get("ONLINE COURSES") or []
        for cl in cert_lines:
            cleaned = cl.lstrip("-*•· ").strip()
            if cleaned:
                data["certifications"].append(cleaned)

        return data

    def render_official_ats_docx(self, candidate_data: dict, user_id: int, version: int, template_id: str = "minimal_ats") -> str:
        """Render candidate resume JSON into official, ATS-friendly Word document following Abhishek Sharma structure."""
        style_config = TEMPLATE_STYLES.get(template_id, TEMPLATE_STYLES["minimal_ats"])
        font_name = style_config["font_family"]
        accent_color = style_config["accent_color"]
        header_align = style_config["header_align"]

        doc = self._setup_document(template_id=template_id)
        from docx.enum.text import WD_TAB_ALIGNMENT

        contact = candidate_data.get("contact") or {}
        name = (contact.get("name") or "Abhishek Sharma").upper()
        subtitle = (candidate_data.get("subtitle") or contact.get("role") or "").upper()
        
        # 1. HEADER (Name, Subtitle, Contact Details)
        header_p = doc.add_paragraph()
        header_p.alignment = header_align
        header_p.paragraph_format.space_before = Pt(0)
        header_p.paragraph_format.space_after = Pt(2)
        
        run_name = header_p.add_run(name)
        run_name.bold = True
        run_name.font.name = font_name
        run_name.font.size = Pt(20)
        run_name.font.color.rgb = accent_color

        if subtitle:
            sub_p = doc.add_paragraph()
            sub_p.alignment = header_align
            sub_p.paragraph_format.space_before = Pt(2)
            sub_p.paragraph_format.space_after = Pt(6)
            run_sub = sub_p.add_run(subtitle)
            run_sub.bold = True
            run_sub.font.name = font_name
            run_sub.font.size = Pt(12)
            run_sub.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

        # Contact line
        email = contact.get("email") or ""
        phone = contact.get("phone") or ""
        location = contact.get("location") or ""
        
        linkedin = contact.get("linkedin") or ""
        if linkedin:
            linkedin = linkedin.replace("https://", "").replace("http://", "")
            if not linkedin.startswith("linkedin.com/"):
                linkedin = f"linkedin.com/in/{linkedin}" if not linkedin.startswith("in/") else f"linkedin.com/{linkedin}"
            
        github = contact.get("github") or ""
        if github:
            github = github.replace("https://", "").replace("http://", "")
            if not github.startswith("github.com/"):
                github = f"github.com/{github}"
            
        contact_info = [
            val for val in [phone, email, location, linkedin, github]
            if val and str(val).strip()
        ]
        contact_line = " | ".join(contact_info)
        if contact_line:
            contact_p = doc.add_paragraph()
            contact_p.alignment = header_align
            contact_p.paragraph_format.space_before = Pt(0)
            contact_p.paragraph_format.space_after = Pt(12)
            run_contact = contact_p.add_run(contact_line)
            run_contact.font.name = font_name
            run_contact.font.size = Pt(9.5)
            run_contact.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

        # 2. PROFILE SUMMARY
        summary = candidate_data.get("summary") or candidate_data.get("professional_summary")
        if summary and str(summary).strip():
            self._add_section_title(doc, "Profile Summary", template_id=template_id)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = 3  # Justify
            add_markdown_paragraph_runs(p, str(summary).strip(), font_name, 10)

        # 3. PROFESSIONAL EXPERIENCE
        experience = candidate_data.get("work_experience") or candidate_data.get("experience")
        if experience and len(experience) > 0:
            self._add_section_title(doc, "Work Experience", template_id=template_id)
            tab_pos = Inches(7.5)
            for exp in experience:
                role = exp.get("role") or exp.get("title") or "Software Engineer"
                company = exp.get("company") or ""
                loc = exp.get("location") or ""
                dates = exp.get("dates") or exp.get("duration") or ""
                if not dates:
                    start = exp.get("start_date") or ""
                    end = exp.get("end_date") or ""
                    dates = f"{start} - {end}" if start and end else (start or end)
                
                # Line 1: Role (Bold) & Dates (Right-aligned using dynamic tab stop)
                p1 = doc.add_paragraph()
                p1.paragraph_format.space_before = Pt(4)
                p1.paragraph_format.space_after = Pt(1.5)
                p1.paragraph_format.line_spacing = 1.0
                p1.paragraph_format.keep_with_next = True
                p1.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
                
                run_role = p1.add_run(role)
                run_role.bold = True
                run_role.font.name = font_name
                run_role.font.size = Pt(10.5)
                
                if dates:
                    p1.add_run("\t")
                    run_dates = p1.add_run(str(dates))
                    run_dates.font.name = font_name
                    run_dates.font.size = Pt(9.5)
                    run_dates.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

                # Line 2: Company Name (Italic) | Location
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(1.5)
                p2.paragraph_format.line_spacing = 1.0
                p2.paragraph_format.keep_with_next = True
                p2.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
                
                run_comp = p2.add_run(company)
                run_comp.italic = True
                run_comp.font.name = font_name
                run_comp.font.size = Pt(10)
                run_comp.font.color.rgb = accent_color
                
                if loc:
                    p2.add_run("\t")
                    run_loc = p2.add_run(loc)
                    run_loc.font.name = font_name
                    run_loc.font.size = Pt(9.5)
                    run_loc.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
                
                bullets = exp.get("bullets") or []
                desc = exp.get("description") or exp.get("responsibilities_achievements") or ""
                tech_used = exp.get("tech_used") or exp.get("technologies") or ""
                
                if not bullets and desc:
                    bullets = [b.strip().lstrip("-*•·").strip() for b in desc.split("\n") if b.strip()]
                    
                for b in bullets:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(1)
                    bp.paragraph_format.line_spacing = 1.0
                    add_markdown_paragraph_runs(bp, b, font_name, 9.5)

                if tech_used:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(1)
                    bp.paragraph_format.line_spacing = 1.0
                    
                    run_label = bp.add_run("Technologies used: ")
                    run_label.bold = True
                    run_label.font.name = font_name
                    run_label.font.size = Pt(9.5)
                    
                    tech_str = ", ".join(tech_used) if isinstance(tech_used, list) else str(tech_used)
                    run_val = bp.add_run(tech_str)
                    run_val.font.name = font_name
                    run_val.font.size = Pt(9.5)

        # 4. SKILLS
        skills = candidate_data.get("technical_skills") or candidate_data.get("skills")
        if skills:
            has_data = False
            if isinstance(skills, dict):
                has_data = any(v for v in skills.values())
            elif isinstance(skills, list):
                has_data = len(skills) > 0
            else:
                has_data = bool(skills)
                
            if has_data:
                self._add_section_title(doc, "Skills", template_id=template_id)
                if isinstance(skills, dict):
                    for cat, items in skills.items():
                        if not items:
                            continue
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(1.5)
                        p.paragraph_format.line_spacing = 1.0
                        
                        run_cat = p.add_run(f"{cat}: ")
                        run_cat.bold = True
                        run_cat.font.name = font_name
                        run_cat.font.size = Pt(9.5)
                        
                        items_str = ", ".join(items) if isinstance(items, list) else str(items)
                        items_str = fix_concatenated_skills(items_str)
                        run_items = p.add_run(items_str)
                        run_items.font.name = font_name
                        run_items.font.size = Pt(9.5)
                elif isinstance(skills, list):
                    if len(skills) > 0 and isinstance(skills[0], dict) and ("category" in skills[0] or "name" in skills[0]):
                        for sk in skills:
                            if not isinstance(sk, dict):
                                continue
                            cat = sk.get("category") or sk.get("name") or ""
                            items = sk.get("items") or sk.get("value") or []
                            if not cat or not items:
                                continue
                            p = doc.add_paragraph()
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(1.5)
                            p.paragraph_format.line_spacing = 1.0
                            
                            run_cat = p.add_run(f"{cat}: ")
                            run_cat.bold = True
                            run_cat.font.name = font_name
                            run_cat.font.size = Pt(9.5)
                            
                            items_str = ", ".join(items) if isinstance(items, list) else str(items)
                            items_str = fix_concatenated_skills(items_str)
                            run_items = p.add_run(items_str)
                            run_items.font.name = font_name
                            run_items.font.size = Pt(9.5)
                    else:
                        flat_items = []
                        for sk in skills:
                            if isinstance(sk, dict):
                                items = sk.get("items") or []
                                if isinstance(items, list):
                                    flat_items.extend([str(x) for x in items])
                                else:
                                    flat_items.append(str(items))
                            else:
                                flat_items.append(str(sk))
                        
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(1.5)
                        p.paragraph_format.line_spacing = 1.0
                        
                        items_str = ", ".join(flat_items)
                        items_str = fix_concatenated_skills(items_str)
                        run_items = p.add_run(items_str)
                        run_items.font.name = font_name
                        run_items.font.size = Pt(9.5)

        # 5. EDUCATION
        education = candidate_data.get("education")
        if education and len(education) > 0:
            self._add_section_title(doc, "Education", template_id=template_id)
            tab_pos = Inches(7.5)
            for edu in education:
                deg = edu.get("degree") or edu.get("specialization") or "Degree"
                if edu.get("specialization") and edu.get("degree") and edu.get("specialization") != edu.get("degree"):
                    deg = f"{edu['degree']} in {edu['specialization']}"
                inst = edu.get("institution") or edu.get("school") or ""
                loc = edu.get("location") or ""
                dates = edu.get("dates") or edu.get("graduation_year") or ""
                if not dates:
                    start = edu.get("start_date") or ""
                    end = edu.get("end_date") or ""
                    dates = f"{start} - {end}" if start and end else (start or end)
                grade = edu.get("grade") or edu.get("gpa") or edu.get("cgpa_percentage") or ""
                
                # Single paragraph line for Education
                p1 = doc.add_paragraph()
                p1.paragraph_format.space_before = Pt(4)
                p1.paragraph_format.space_after = Pt(1.5)
                p1.paragraph_format.line_spacing = 1.0
                p1.paragraph_format.keep_with_next = True
                p1.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
                
                # Left side
                run_inst = p1.add_run(inst)
                run_inst.bold = True
                run_inst.font.name = font_name
                run_inst.font.size = Pt(10)
                
                if deg:
                    p1.add_run(" | ")
                    run_deg = p1.add_run(deg)
                    run_deg.italic = True
                    run_deg.font.name = font_name
                    run_deg.font.size = Pt(10)
                    
                # Right side
                p1.add_run("\t")
                
                right_parts = []
                if dates:
                    right_parts.append(str(dates))
                if grade:
                    grade_str = f"CGPA: {grade}" if not str(grade).lower().startswith(("cgpa", "grade")) else str(grade)
                    right_parts.append(grade_str)
                    
                right_line = " | ".join(right_parts)
                if right_line:
                    run_right = p1.add_run(right_line)
                    run_right.italic = True
                    run_right.font.name = font_name
                    run_right.font.size = Pt(10)
                    run_right.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

        # 6. PROJECTS
        projects = candidate_data.get("projects")
        if projects and len(projects) > 0:
            self._add_section_title(doc, "Projects", template_id=template_id)
            tab_pos = Inches(7.5)
            for proj in projects:
                name = proj.get("name") or proj.get("title") or "Project"
                tech = proj.get("technologies") or proj.get("tech_used") or ""
                dates = proj.get("dates") or proj.get("date") or ""
                link = proj.get("link") or proj.get("url") or proj.get("demo_link") or ""
                
                # Single paragraph line for Project Header
                p1 = doc.add_paragraph()
                p1.paragraph_format.space_before = Pt(4)
                p1.paragraph_format.space_after = Pt(1.5)
                p1.paragraph_format.line_spacing = 1.0
                p1.paragraph_format.keep_with_next = True
                p1.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
                
                # Left side: **Name** | *Technologies*
                run_name = p1.add_run(name)
                run_name.bold = True
                run_name.font.name = font_name
                run_name.font.size = Pt(10)
                
                if tech:
                    p1.add_run(" | ")
                    tech_clean = ", ".join(tech) if isinstance(tech, list) else str(tech)
                    tech_clean = fix_concatenated_skills(tech_clean)
                    run_tech = p1.add_run(tech_clean)
                    run_tech.italic = True
                    run_tech.font.name = font_name
                    run_tech.font.size = Pt(10)
                    
                # Right side: *Dates* | *Link*
                p1.add_run("\t")
                
                right_parts = []
                if dates:
                    right_parts.append(str(dates))
                if link:
                    right_parts.append(link)
                    
                right_line = " | ".join(right_parts)
                if right_line:
                    run_right = p1.add_run(right_line)
                    run_right.italic = True
                    run_right.font.name = font_name
                    run_right.font.size = Pt(10)
                    run_right.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
                
                bullets = proj.get("bullets") or []
                desc = proj.get("description") or proj.get("contribution_impact") or ""
                if not bullets and desc:
                    bullets = [b.strip().lstrip("-*•·").strip() for b in desc.split("\n") if b.strip()]
                    
                for b in bullets:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(1)
                    bp.paragraph_format.line_spacing = 1.0
                    add_markdown_paragraph_runs(bp, b, font_name, 9.5)

        # 7. ONLINE COURSES & CERTIFICATIONS
        certs = candidate_data.get("certifications")
        if certs and len(certs) > 0:
            self._add_section_title(doc, "Online Courses & Certifications", template_id=template_id)
            for c in certs:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.line_spacing = 1.0
                
                if isinstance(c, dict):
                    name = c.get("name") or c.get("title") or ""
                    issuer = c.get("issuer") or c.get("authority") or c.get("institution") or ""
                    date = c.get("date") or c.get("year") or ""
                    link = c.get("credential_link") or c.get("link") or c.get("url") or ""
                    
                    run_name = bp.add_run(name)
                    run_name.bold = True
                    run_name.font.name = font_name
                    run_name.font.size = Pt(9.5)
                    
                    rest_str = ""
                    if issuer:
                        rest_str += f" — {issuer}"
                    if date:
                        rest_str += f" ({date})"
                    if rest_str:
                        run_rest = bp.add_run(rest_str)
                        run_rest.font.name = font_name
                        run_rest.font.size = Pt(9.5)
                        
                    if link:
                        run_sep = bp.add_run(" | ")
                        run_sep.font.name = font_name
                        run_sep.font.size = Pt(9.5)
                        
                        run_link = bp.add_run(link)
                        run_link.font.name = font_name
                        run_link.font.size = Pt(9.5)
                        run_link.font.color.rgb = accent_color
                        run_link.underline = True
                else:
                    run_c = bp.add_run(str(c))
                    run_c.font.name = font_name
                    run_c.font.size = Pt(9.5)

        # 8. ACHIEVEMENTS & EXTRACURRICULAR
        ach = candidate_data.get("achievements")
        if ach and len(ach) > 0:
            self._add_section_title(doc, "Achievements & Extracurricular", template_id=template_id)
            for a in ach:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.line_spacing = 1.0
                
                if isinstance(a, dict):
                    text_val = a.get("description") or a.get("title") or a.get("name") or a.get("details") or str(a)
                else:
                    text_val = str(a)
                    
                add_markdown_paragraph_runs(bp, text_val, font_name, 9.5)

        self._enforce_tight_spacing(doc, template_id=template_id)
        output_filename = f"final_documents/user_{user_id}/resume_v{version}.docx"
        output_path = WORKSPACE_ROOT / output_filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path.resolve()))
        return str(output_path.resolve())

    def compile_docx_state(self, state: dict[str, Any], output_path: str | Path) -> Path:
        """Compile a state JSON resume into a professional Word docx file."""
        resume = state.get("optimized_resume_json") or state.get("user_resume_json") or state.get("extracted_facts") or {}
        raw_text = state.get("optimized_resume") or state.get("resume_text") or ""
        
        # Determine if we should parse raw text instead of dummy dictionary
        if isinstance(resume, str):
            resume = self.parse_resume_text_to_dict(resume)
        elif isinstance(resume, dict):
            exp = resume.get("experience") or []
            if len(exp) == 1 and "EDUCATION" in (exp[0].get("description") or ""):
                resume = self.parse_resume_text_to_dict(exp[0].get("description"))
            elif not resume or (not resume.get("experience") and not resume.get("projects") and raw_text):
                resume = self.parse_resume_text_to_dict(raw_text)
        else:
            resume = self.parse_resume_text_to_dict(raw_text)
                
        user_id = state.get("user_id") or 1
        attempt = state.get("attempt_count") or 1
        template_id = state.get("template_id") or "minimal_ats"
        
        abs_path_str = self.render_official_ats_docx(resume, user_id, attempt, template_id=template_id)
        dest_path = self._workspace_path(output_path)
        dest_path.parent.mkdir(parents=True, exist_ok=True)
        if str(dest_path.resolve()) != abs_path_str:
            import shutil
            shutil.copy2(abs_path_str, str(dest_path.resolve()))
        return dest_path

    def _render_docx_template(self, doc: docx.Document, candidate_data: dict[str, Any], template_id: str) -> None:
        """Replace placeholders and dynamically expand lists/tables inside the loaded template document."""
        style_config = TEMPLATE_STYLES.get(template_id, TEMPLATE_STYLES["minimal_ats"])
        font_name = style_config["font_family"]
        accent_color = style_config["accent_color"]
        header_align = style_config["header_align"]

        contact = candidate_data.get("contact") or {}
        replacements = {
            "{{name}}": (contact.get("name") or "Abhishek Sharma").upper(),
            "{{subtitle}}": (candidate_data.get("subtitle") or contact.get("role") or "").upper(),
            "{{phone}}": contact.get("phone") or "",
            "{{email}}": contact.get("email") or "",
            "{{location}}": contact.get("location") or "",
            "{{linkedin}}": contact.get("linkedin") or "",
            "{{github}}": contact.get("github") or "",
            "{{portfolio}}": contact.get("portfolio") or contact.get("website") or "",
            "{{summary}}": candidate_data.get("summary") or "",
        }

        # 1. Reconstruct contact details line dynamically from dictionary values
        email = contact.get("email") or ""
        phone = contact.get("phone") or ""
        location = contact.get("location") or ""
        
        linkedin = contact.get("linkedin") or ""
        if linkedin:
            linkedin = linkedin.replace("https://", "").replace("http://", "").replace("linkedin.com/in/", "").replace("linkedin.com/", "")
            linkedin = f"in/{linkedin}"
            
        github = contact.get("github") or ""
        if github:
            github = github.replace("https://", "").replace("http://", "").replace("github.com/", "")
            github = f"github.com/{github}"
            
        contact_info = [
            val for val in [email, phone, location, linkedin, github]
            if val and str(val).strip()
        ]
        contact_line = " | ".join(contact_info)

        # 2. Build separated name, subtitle/role, and contact paragraphs dynamically
        for p in list(doc.paragraphs):
            if "{{name}}" in p.text:
                # Insert Name
                p_name = p.insert_paragraph_before()
                p_name.alignment = header_align
                p_name.paragraph_format.space_before = Pt(0)
                p_name.paragraph_format.space_after = Pt(2)
                r_name = p_name.add_run(replacements["{{name}}"])
                r_name.bold = True
                r_name.font.name = font_name
                r_name.font.size = Pt(18)
                r_name.font.color.rgb = accent_color

                # Insert Subtitle
                subtitle_val = replacements["{{subtitle}}"]
                if subtitle_val:
                    p_sub = p.insert_paragraph_before()
                    p_sub.alignment = header_align
                    p_sub.paragraph_format.space_before = Pt(2)
                    p_sub.paragraph_format.space_after = Pt(6)
                    r_sub = p_sub.add_run(subtitle_val)
                    r_sub.bold = True
                    r_sub.font.name = font_name
                    r_sub.font.size = Pt(12)
                    r_sub.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

                # Insert Contact details
                if contact_line:
                    p_contact = p.insert_paragraph_before()
                    p_contact.alignment = header_align
                    p_contact.paragraph_format.space_before = Pt(2)
                    p_contact.paragraph_format.space_after = Pt(12)
                    r_contact = p_contact.add_run(contact_line)
                    r_contact.font.name = font_name
                    r_contact.font.size = Pt(9.5)
                    r_contact.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

                # Remove the original concatenated placeholder paragraph
                p._p.getparent().remove(p._p)
                break

        # 3. Replace any other remaining basic placeholders in standard text blocks (e.g. {{summary}})
        for p in doc.paragraphs:
            text = p.text
            for k, v in replacements.items():
                if k in text:
                    for run in p.runs:
                        if k in run.text:
                            run.text = run.text.replace(k, v)
                            run.font.name = font_name
                    if k in p.text:
                        p.text = p.text.replace(k, v)

        from docx.enum.text import WD_TAB_ALIGNMENT
        for p in list(doc.paragraphs):
            text = p.text.strip()
            if text == "{{experience}}":
                p.text = ""
                experience = candidate_data.get("experience") or []
                margin_val = style_config.get("margins", 0.75)
                tab_pos = Inches(8.5 - (2 * margin_val))
                for exp in experience:
                    role = exp.get("role") or "Software Engineer"
                    company = exp.get("company") or ""
                    loc = exp.get("location") or ""
                    dates = exp.get("dates") or exp.get("start_date") or ""
                    if not dates:
                        start = exp.get("start_date") or ""
                        end = exp.get("end_date") or ""
                        dates = f"{start} - {end}" if start and end else (start or end)
                    
                    # Line 1: Role (Bold) & Dates (Right-aligned using dynamic tab stop)
                    p1 = p.insert_paragraph_before()
                    p1.paragraph_format.space_before = Pt(6)
                    p1.paragraph_format.space_after = Pt(0)
                    p1.paragraph_format.keep_with_next = True
                    p1.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
                    
                    r_role = p1.add_run(role)
                    r_role.bold = True
                    r_role.font.name = font_name
                    r_role.font.size = Pt(10.5)
                    
                    if dates:
                        p1.add_run("\t")
                        r_dates = p1.add_run(dates)
                        r_dates.font.name = font_name
                        r_dates.font.size = Pt(9.5)
                        r_dates.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

                    # Line 2: Company | Location
                    p2 = p.insert_paragraph_before()
                    p2.paragraph_format.space_before = Pt(0)
                    p2.paragraph_format.space_after = Pt(3)
                    p2.paragraph_format.keep_with_next = True
                    
                    r_comp = p2.add_run(company)
                    r_comp.italic = True
                    r_comp.font.name = font_name
                    r_comp.font.size = Pt(10)
                    r_comp.font.color.rgb = accent_color
                    
                    if loc:
                        p2.add_run("  |  ")
                        r_loc = p2.add_run(loc)
                        r_loc.font.name = font_name
                        r_loc.font.size = Pt(9.5)
                        r_loc.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                    
                    bullets = exp.get("bullets") or []
                    desc = exp.get("description") or ""
                    if not bullets and desc:
                        bullets = [b.strip().lstrip("-*•·").strip() for b in desc.split("\n") if b.strip()]
                        
                    for b in bullets:
                        bp = p.insert_paragraph_before(style="List Bullet")
                        bp.paragraph_format.space_before = Pt(0)
                        bp.paragraph_format.space_after = Pt(2)
                        
                        m = re.match(r'^(\*\*[^*]+\*\*)(.*)$', b)
                        if m:
                            verb = m.group(1).replace("**", "")
                            rest = m.group(2)
                            rv = bp.add_run(verb)
                            rv.bold = True
                            rv.font.name = font_name
                            rv.font.size = Pt(9.5)
                            rr = bp.add_run(rest)
                            rr.font.name = font_name
                            rr.font.size = Pt(9.5)
                        else:
                            words = b.split(" ", 1)
                            if len(words) > 1 and re.match(r'^[A-Z][a-z]+$', words[0]):
                                rv = bp.add_run(words[0] + " ")
                                rv.bold = True
                                rv.font.name = font_name
                                rv.font.size = Pt(9.5)
                                rr = bp.add_run(words[1])
                                rr.font.name = font_name
                                rr.font.size = Pt(9.5)
                            else:
                                run_b = bp.add_run(b)
                                run_b.font.name = font_name
                                run_b.font.size = Pt(9.5)
                p._p.getparent().remove(p._p)

            elif text == "{{skills}}":
                p.text = ""
                skills = candidate_data.get("skills") or []
                for sk in skills:
                    bp = p.insert_paragraph_before(style="List Bullet")
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(2)
                    
                    if isinstance(sk, dict) and "category" in sk:
                        cat = sk["category"]
                        items = ", ".join(sk["items"]) if isinstance(sk["items"], list) else str(sk["items"])
                        items = fix_concatenated_skills(items)
                        
                        run_cat = bp.add_run(f"{cat}: ")
                        run_cat.bold = True
                        run_cat.font.name = font_name
                        run_cat.font.size = Pt(9.5)
                        
                        run_items = bp.add_run(items)
                        run_items.font.name = font_name
                        run_items.font.size = Pt(9.5)
                    else:
                        sk_str = fix_concatenated_skills(str(sk))
                        run_sk = bp.add_run(sk_str)
                        run_sk.font.name = font_name
                        run_sk.font.size = Pt(9.5)
                p._p.getparent().remove(p._p)

            elif text == "{{education}}":
                p.text = ""
                education = candidate_data.get("education") or []
                margin_val = style_config.get("margins", 0.75)
                tab_pos = Inches(8.5 - (2 * margin_val))
                for edu in education:
                    deg = edu.get("degree") or "Degree"
                    inst = edu.get("institution") or edu.get("school") or ""
                    loc = edu.get("location") or ""
                    dates = edu.get("dates") or edu.get("start_date") or ""
                    if not dates:
                        start = edu.get("start_date") or ""
                        end = edu.get("end_date") or ""
                        dates = f"{start} - {end}" if start and end else (start or end)
                    grade = edu.get("grade") or edu.get("gpa") or ""
                    
                    # Line 1: Degree & Dates (Right-aligned using dynamic tab stop)
                    p1 = p.insert_paragraph_before()
                    p1.paragraph_format.space_before = Pt(4)
                    p1.paragraph_format.space_after = Pt(0)
                    p1.paragraph_format.keep_with_next = True
                    p1.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
                    
                    run_deg = p1.add_run(deg)
                    run_deg.bold = True
                    run_deg.font.name = font_name
                    run_deg.font.size = Pt(10.5)
                    
                    if dates:
                        p1.add_run("\t")
                        run_dates = p1.add_run(dates)
                        run_dates.font.name = font_name
                        run_dates.font.size = Pt(9.5)
                        run_dates.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                    
                    # Line 2: Institution & Location | Grade (if present)
                    p2 = p.insert_paragraph_before()
                    p2.paragraph_format.space_before = Pt(0)
                    p2.paragraph_format.space_after = Pt(4)
                    p2.paragraph_format.keep_with_next = True
                    
                    if inst:
                        inst_loc = f"{inst}, {loc}" if loc else inst
                        run_inst = p2.add_run(inst_loc)
                        run_inst.italic = True
                        run_inst.font.name = font_name
                        run_inst.font.size = Pt(10)
                        run_inst.font.color.rgb = accent_color
                        
                    if grade:
                        p2.add_run("  |  ")
                        run_grade = p2.add_run(grade)
                        run_grade.bold = True
                        run_grade.font.name = font_name
                        run_grade.font.size = Pt(9.5)
                        run_grade.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                p._p.getparent().remove(p._p)

            elif text == "{{projects}}":
                p.text = ""
                projects = candidate_data.get("projects") or []
                margin_val = style_config.get("margins", 0.75)
                tab_pos = Inches(8.5 - (2 * margin_val))
                for proj in projects:
                    name = proj.get("name") or "Project"
                    tech = proj.get("technologies") or ""
                    dates = proj.get("dates") or ""
                    link = proj.get("link") or proj.get("url") or ""
                    
                    # Line 1: Project Name & Dates (Right-aligned using dynamic tab stop)
                    p1 = p.insert_paragraph_before()
                    p1.paragraph_format.space_before = Pt(4)
                    p1.paragraph_format.space_after = Pt(0)
                    p1.paragraph_format.keep_with_next = True
                    p1.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
                    
                    run_name = p1.add_run(name)
                    run_name.bold = True
                    run_name.font.name = font_name
                    run_name.font.size = Pt(10.5)
                    
                    if dates:
                        p1.add_run("\t")
                        run_dates = p1.add_run(dates)
                        run_dates.font.name = font_name
                        run_dates.font.size = Pt(9.5)
                        run_dates.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                    
                    # Line 2: Technologies | Link
                    p2 = p.insert_paragraph_before()
                    p2.paragraph_format.space_before = Pt(0)
                    p2.paragraph_format.space_after = Pt(3)
                    p2.paragraph_format.keep_with_next = True
                    
                    if tech:
                        tech_clean = fix_concatenated_skills(tech)
                        run_tech = p2.add_run(tech_clean)
                        run_tech.italic = True
                        run_tech.font.name = font_name
                        run_tech.font.size = Pt(9.5)
                        run_tech.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                        
                    if link:
                        if tech:
                            p2.add_run("  |  ")
                        run_link = p2.add_run("Link")
                        run_link.underline = True
                        run_link.font.name = font_name
                        run_link.font.size = Pt(10)
                        run_link.font.color.rgb = accent_color
                    
                    bullets = proj.get("bullets") or []
                    desc = proj.get("description") or ""
                    if not bullets and desc:
                        bullets = [b.strip().lstrip("-*•·").strip() for b in desc.split("\n") if b.strip()]
                        
                    for b in bullets:
                        bp = p.insert_paragraph_before(style="List Bullet")
                        bp.paragraph_format.space_before = Pt(0)
                        bp.paragraph_format.space_after = Pt(2)
                        
                        m = re.match(r'^(\*\*[^*]+\*\*)(.*)$', b)
                        if m:
                            verb = m.group(1).replace("**", "")
                            rest = m.group(2)
                            rv = bp.add_run(verb)
                            rv.bold = True
                            rv.font.name = font_name
                            rv.font.size = Pt(9.5)
                            rr = bp.add_run(rest)
                            rr.font.name = font_name
                            rr.font.size = Pt(9.5)
                        else:
                            words = b.split(" ", 1)
                            if len(words) > 1 and re.match(r'^[A-Z][a-z]+$', words[0]):
                                rv = bp.add_run(words[0] + " ")
                                rv.bold = True
                                rv.font.name = font_name
                                rv.font.size = Pt(9.5)
                                rr = bp.add_run(words[1])
                                rr.font.name = font_name
                                rr.font.size = Pt(9.5)
                            else:
                                run_b = bp.add_run(b)
                                run_b.font.name = font_name
                                run_b.font.size = Pt(9.5)
                p._p.getparent().remove(p._p)
        self._enforce_tight_spacing(doc, template_id=template_id)

    def compile_agent_state(self, state: AgentState | dict[str, Any], document_type: str = "resume") -> Path:
        """Render and save a state document as docx under storage_workspace/final_documents or storage_workspace/generated."""
        payload = state.model_dump() if hasattr(state, "model_dump") else state
        user_id = payload.get("user_id") or 1
        attempt = payload.get("attempt_count") or 1

        if document_type == "resume":
            filename = f"final_documents/user_{user_id}/resume_v{attempt}.docx"
            return self.compile_docx_state(payload, filename)
        else:
            session = str(payload.get("session_id") or "unsessioned").replace("/", "_").replace("\\", "_")
            filename = f"generated/{document_type}_{session}.docx"
            return self.compile_docx_cover_letter(payload, filename)

    def compile_docx_cover_letter(self, state: dict[str, Any], output_path: str | Path) -> Path:
        """Compile cover letter text from state into a Word docx file."""
        doc = docx.Document()
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        name = state.get("metadata", {}).get("candidate_name") or "Candidate"
        cover_letter_text = state.get("cover_letter", "") or state.get("cover_letter_text", "")

        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(17, 24, 39)

        body_p = doc.add_paragraph()
        body_p.paragraph_format.space_before = Pt(12)
        run_body = body_p.add_run(cover_letter_text)
        run_body.font.name = 'Arial'
        run_body.font.size = Pt(11)

        self._enforce_tight_spacing(doc, template_id="minimal_ats")
        path = self._workspace_path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return path

    @staticmethod
    def _workspace_path(output_path: str | Path) -> Path:
        root = WORKSPACE_ROOT.resolve()
        path = Path(output_path)
        resolved = (root / path).resolve() if not path.is_absolute() else path.resolve()
        if root not in resolved.parents and resolved != root:
            resolved = (root / path.name).resolve()
        return resolved
