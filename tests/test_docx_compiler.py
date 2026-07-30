# File: tests/test_docx_compiler.py
"""Tests for Word Document compiler logic."""

from __future__ import annotations

from pathlib import Path
from src.utils.docx_compiler import DocxCompiler


def test_render_official_ats_docx() -> None:
    compiler = DocxCompiler()
    candidate_data = {
        "contact": {
            "name": "Test User",
            "email": "test@example.com",
            "phone": "123-456-7890",
            "location": "San Francisco, CA",
            "github": "github.com/test",
            "linkedin": "linkedin.com/in/test"
        },
        "summary": "Experienced testing engineer",
        "skills": ["Python", "Pytest", "Word"],
        "experience": [
            {
                "company": "QA Corp",
                "role": "Test Lead",
                "start_date": "2020",
                "end_date": "2023",
                "description": "- Tested code.\n- Wrote scripts."
            }
        ],
        "projects": [
            {
                "name": "Test Tool",
                "description": "- Automated testing.\n- Created reports."
            }
        ],
        "education": [
            {
                "institution": "State University",
                "degree": "B.S. Computer Science",
                "start_date": "2016",
                "end_date": "2020"
            }
        ]
    }
    
    doc_path = compiler.render_official_ats_docx(candidate_data, user_id=99, version=1)
    path = Path(doc_path)
    assert path.exists()
    assert "user_99" in str(path)
    assert "resume_v1.docx" in str(path)


def test_compile_agent_state() -> None:
    compiler = DocxCompiler()
    state = {
        "user_id": 42,
        "attempt_count": 3,
        "optimized_resume_json": {
            "contact": {"name": "State User", "email": "state@example.com"},
            "summary": "Expert dev",
            "skills": ["Python"],
            "experience": []
        }
    }
    docx_path = compiler.compile_agent_state(state, "resume")
    assert "final_documents" in str(docx_path)
    assert "user_42" in str(docx_path)
    assert "resume_v3.docx" in str(docx_path)
    assert docx_path.exists()


def test_compile_docx_cover_letter() -> None:
    compiler = DocxCompiler()
    state = {
        "metadata": {"candidate_name": "Cover Letter User"},
        "cover_letter": "Dear Recruiter, this is a test cover letter.",
        "session_id": "test_session_123"
    }
    docx_path = compiler.compile_agent_state(state, "cover_letter")
    assert "generated" in str(docx_path)
    assert "cover_letter_test_session_123.docx" in str(docx_path)
    assert docx_path.exists()


def test_parse_resume_text_to_dict() -> None:
    compiler = DocxCompiler()
    raw_text = (
        "Abhishek Sharma\n"
        "AI/ML & Python Developer\n"
        "shabhishek055@gmail.com | +91 9575676062 | Indore, India | linkedin.com/in/Abhishek\n\n"
        "---\n"
        "PROFILE SUMMARY\n"
        "---\n"
        "Results-driven AI/ML Engineer with a proven track record.\n\n"
        "---\n"
        "EDUCATION\n"
        "---\n"
        "Bachelor of Technology (IT Engineering) | Rajiv Gandhi Prodhogiki Vishvvidhyalya, Bhopal | Indore | 2020-2024 | 7.6 CGPA\n\n"
        "---\n"
        "PROJECTS\n"
        "---\n"
        "TechnikalPAI - A Behavior-Based Friendly AI | Open-AI, Flask | 04/2024 - present\n"
        "• Developed behavior-based Friendly AI.\n"
        "• Achieved 85% accuracy.\n\n"
        "---\n"
        "SKILLS\n"
        "---\n"
        "Languages: Python, SQL, R\n"
        "ML Frameworks: TensorFlow, PyTorch"
    )
    
    parsed = compiler.parse_resume_text_to_dict(raw_text)
    assert parsed["contact"]["name"] == "Abhishek Sharma"
    assert parsed["subtitle"] == "AI/ML & Python Developer"
    assert parsed["contact"]["email"] == "shabhishek055@gmail.com"
    assert parsed["contact"]["phone"] == "+91 9575676062"
    assert len(parsed["education"]) == 1
    assert parsed["education"][0]["degree"] == "Bachelor of Technology (IT Engineering)"
    assert len(parsed["projects"]) == 1
    assert parsed["projects"][0]["name"] == "TechnikalPAI - A Behavior-Based Friendly AI"
    assert len(parsed["projects"][0]["bullets"]) == 2
    assert len(parsed["skills"]) == 2
    assert parsed["skills"][0]["category"] == "Languages"


def test_docx_compiler_presets() -> None:
    import docx
    compiler = DocxCompiler()
    candidate_data = {
        "contact": {
            "name": "Preset User",
            "email": "preset@example.com"
        },
        "summary": "Testing presets",
        "skills": ["Python"]
    }
    
    # 1. Test minimal_ats
    p1 = compiler.render_official_ats_docx(candidate_data, user_id=1, version=1, template_id="minimal_ats")
    doc1 = docx.Document(p1)
    assert doc1.sections[0].top_margin.inches == 0.75
    assert doc1.styles['Normal'].font.name == 'Arial'

    # 2. Test modern_tech
    p2 = compiler.render_official_ats_docx(candidate_data, user_id=1, version=2, template_id="modern_tech")
    doc2 = docx.Document(p2)
    assert doc2.sections[0].top_margin.inches == 0.6
    assert doc2.styles['Normal'].font.name == 'Calibri'

    # 3. Test classic_executive
    p3 = compiler.render_official_ats_docx(candidate_data, user_id=1, version=3, template_id="classic_executive")
    doc3 = docx.Document(p3)
    assert doc3.sections[0].top_margin.inches == 0.75
    assert doc3.styles['Normal'].font.name == 'Georgia'

    # 4. Test compact_onepage
    p4 = compiler.render_official_ats_docx(candidate_data, user_id=1, version=4, template_id="compact_onepage")
    doc4 = docx.Document(p4)
    assert doc4.sections[0].top_margin.inches == 0.5
    assert doc4.styles['Normal'].font.name == 'Arial'


def test_fix_concatenated_skills() -> None:
    from src.utils.docx_compiler import fix_concatenated_skills
    
    assert fix_concatenated_skills("Go (Golang)Python SQL") == "Go (Golang), Python, SQL"
    assert fix_concatenated_skills("PostgreSQLRedis MongoDB") == "PostgreSQL, Redis, MongoDB"
    assert fix_concatenated_skills("KubernetesTerraformCI/CD") == "Kubernetes, Terraform, CI/CD"
    assert fix_concatenated_skills("Event-DrivenKafka System Design") == "Event-Driven Architecture, Kafka, System Design"
    assert fix_concatenated_skills("DockerKubernetes") == "Docker, Kubernetes"


def test_layout_separation() -> None:
    import docx
    compiler = DocxCompiler()
    candidate_data = {
        "contact": {
            "name": "Alex Mercer",
            "email": "alex@example.com"
        },
        "subtitle": "Senior Backend Engineer",
        "experience": [
            {
                "role": "Lead Architect",
                "company": "CloudScale Technologies",
                "dates": "Jan 2023 – Present",
                "location": "Remote"
            }
        ]
    }
    
    doc_path = compiler.render_official_ats_docx(candidate_data, user_id=22, version=1, template_id="minimal_ats")
    doc = docx.Document(doc_path)
    
    # Assert header name & role separation and uppercase
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    assert "ALEX MERCER" in paragraphs
    assert "SENIOR BACKEND ENGINEER" in paragraphs
    assert "ALEX MERCERSENIOR BACKEND ENGINEER" not in paragraphs
    
    # Verify name is 18pt and role is 12pt
    name_p = next(p for p in doc.paragraphs if p.text == "ALEX MERCER")
    assert name_p.runs[0].font.size == docx.shared.Pt(18)
    assert name_p.runs[0].bold is True
    
    role_p = next(p for p in doc.paragraphs if p.text == "SENIOR BACKEND ENGINEER")
    assert role_p.runs[0].font.size == docx.shared.Pt(12)
    assert role_p.runs[0].bold is True
    
    # Assert Role and Dates are on Line 1, Company & Location on Line 2
    assert any("Lead Architect" in p for p in paragraphs)
    role_date_p = next(p for p in doc.paragraphs if "Lead Architect" in p.text)
    assert "Jan 2023 – Present" in role_date_p.text
    assert "CloudScale Technologies" not in role_date_p.text
    
    comp_loc_p = next(p for p in doc.paragraphs if "CloudScale Technologies" in p.text)
    assert "Remote" in comp_loc_p.text
    assert "Lead Architect" not in comp_loc_p.text
    
    # Assert contact line is sanitized
    contact_p = next(p for p in doc.paragraphs if "alex@example.com" in p.text)
    assert contact_p.text == "alex@example.com"
    
    # Assert paragraph spacing on bullets is 1.5 Pt
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    for b in bullets:
        assert b.paragraph_format.space_before == docx.shared.Pt(0)
        assert b.paragraph_format.space_after == docx.shared.Pt(1.5)
        assert b.paragraph_format.line_spacing == 1.15



